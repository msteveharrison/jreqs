#!/usr/bin/env python3

import argparse 
import docx
import json
import reqgraph
import glob
import copy

# Helper function to insert an item at a location
# other than the end of the document
def move_after(first, second):
    e1, e2 = first._element, second._element
    e1.addnext(e2)

# Programmatically remove everything in a section.
def remove_content(anchor):

    # Move to the next element right away. 
    # We dont want to delete the anchor element.
    el = anchor._element.getnext()

    while True:

        # Remove everything until we encounter a non-empty paragraph
        if isinstance(el, docx.oxml.text.paragraph.CT_P):
            para = docx.text.paragraph.Paragraph(el, el.getparent())
            if para.text != "":
                return

        # Get the next element before we delete this one
        try:
            n = el.getnext()
            el.getparent().remove(el)
            el = n
        except AttributeError:
            return # End of file.
    


if __name__ == '__main__':

    parser = argparse.ArgumentParser(description='Populate DOCX file with requirements derived from JSON files.')
    parser.add_argument('json', help='JSON file to parse.')
    parser.add_argument('din', help='MS-Word doc to read in.')
    parser.add_argument('dout', help='MS-Word doc to write out.')
    args = parser.parse_args()

    # Try to open the input document.
    try:
        din = docx.Document(args.din)
    except docx.opc.exceptions.PackageNotFoundError as e:
        print(e)
        exit(1)

    # Try to open the output document.
    try:
        jsonfile = open(args.json, 'r')
    except FileNotFoundError as e:
        print(e)
        exit(1)

    # Try to load the json
    reqs = json.load(jsonfile)

    # Try to load the dependencies
    g = reqgraph.build(glob.glob('*.json'))

    # TODO: Validate JSON

    # Try to find the heading to insert at.
    # A heading in Word is just a paragraph with a style applied.
    anchor = None
    header = reqs['group']
    for p in din.paragraphs:
        if header == p.text:
            anchor = p

    if anchor is None:
        print('Could not find header: %s' %(header))
        exit(1)

    remove_content(anchor)  

    # Add the tables. 
    for req in reqs['reqs']:

        tbl = din.add_table(0, 2, style='Table Grid')
        tbl.autofit = True
        tbl.columns[0].width = 914400
        tbl.columns[1].width = int(5.5*914400)

        row = tbl.add_row()
        
        row.cells[0].text = req['num']
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = req['dsc']
        row.cells[1].paragraphs[0].runs[0].bold = True

        row = tbl.add_row()
        row.cells[0].text = 'Requirement'   
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = req['txt']

        row = tbl.add_row()
        row.cells[0].text = 'Rationale'
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = req['rat']

        refs = g[req['num']].refines()
        if len(refs) > 0:
            row = tbl.add_row()
            row.cells[0].text = 'Refines'
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = 'TODO'

        refbys = g[req['num']].refinedBy()
        if len(refbys) > 0:
            row = tbl.add_row()
            row.cells[0].text = 'Refined By'
            row.cells[0].paragraphs[0].runs[0].bold = True      
            row.cells[1].text = '\n'.join([x.num() + " (" + x.dsc() + ")" for x in refbys])

        row = tbl.add_row()
        row.cells[0].text = 'Verification'
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = req['ver']

        move_after(anchor, tbl)

        anchor = din.add_paragraph('')
        move_after(tbl, anchor)

    try:
        din.save(args.dout)
    except OSError as e:
        print(e)
        exit(1)

    exit(0)

